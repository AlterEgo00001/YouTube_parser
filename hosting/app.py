import cv2
import youtube_dl
import pytesseract
from PIL import Image
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from flask import Flask, request, send_file

app = Flask(__name__)

def download_video(url):
    ydl_opts = {'format': 'bestvideo+bestaudio/best', 'outtmpl': 'video.mp4'}
    with youtube_dl.YoutubeDL(ydl_opts) as ydl:
        ydl.download([url])

def extract_frames(video_path, frame_rate=30):
    cap = cv2.VideoCapture(video_path)
    i = 0
    while cap.isOpened():
        ret, frame = cap.read()
        if not ret:
            break
        if i % frame_rate == 0:
            cv2.imwrite(f'frame_{i}.png', frame)
        i += 1
    cap.release()
    cv2.destroyAllWindows()

def extract_text_from_image(image_path):
    img = Image.open(image_path)
    gray = img.convert('L')
    bw = gray.point(lambda x: 0, '1')
    text = pytesseract.image_to_string(bw)
    return text

def create_pdf(images_with_text, output_path):
    c = canvas.Canvas(output_path, pagesize=letter)
    width, height = letter
    for image_path, text in images_with_text:
        c.drawImage(image_path, 0, height/2, width=width, height=height/2)
        c.drawString(100, height/2 - 50, text)
        c.showPage()
    c.save()

def create_doc(images_with_text, output_path):
    doc = Document()
    for image_path, text in images_with_text:
        doc.add_picture(image_path, width=Inches(6))
        doc.add_paragraph(text)
    doc.save(output_path)

@app.route('/parse', methods=['POST'])
def parse_video():
    url = request.form['url']
    output_format = request.form['format']
    download_video(url)
    extract_frames('video.mp4', frame_rate=30)
    images_with_text = [(f'frame_{i}.png', extract_text_from_image(f'frame_{i}.png')) for i in range(0, 100, 30)]
    if output_format == 'pdf':
        create_pdf(images_with_text, 'output.pdf')
        return send_file('output.pdf', as_attachment=True)
    elif output_format == 'doc':
        create_doc(images_with_text, 'output.docx')
        return send_file('output.docx', as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
